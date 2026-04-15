"""
File extraction service for PDF, DOCX, and TXT files.
Extracts text content from uploaded files for AI agent processing.
"""
import logging
from typing import Optional
from pathlib import Path
import mimetypes
from pypdf import PdfReader
from docx import Document
from io import BytesIO
from PIL import Image

logger = logging.getLogger(__name__)


class FileExtractor:
    """Extract text content from uploaded files."""
    
    # Supported file extensions
    SUPPORTED_EXTENSIONS = {
        '.pdf': 'pdf',
        '.docx': 'docx',
        '.txt': 'txt',
    }
    
    MAX_FILE_SIZE = 6 * 1024 * 1024  # 6MB
    
    def __init__(self):
        # Initialize mimetypes
        mimetypes.init()
    
    def extract_text_and_images(self, file_content: bytes, filename: str) -> dict:
        """
        Extract text and images from file content.
        Args:
            file_content: Raw file bytes
            filename: Original filename
        Returns:
            dict with keys:
                - text: Extracted text content
                - images: List of extracted images (for PDFs)
                - file_type: Detected file type
                - error: Error message if extraction failed
        """
        try:
            if len(file_content) > self.MAX_FILE_SIZE:
                return {
                    'text': None,
                    'images': [],
                    'file_type': None,
                    'error': f'File size exceeds {self.MAX_FILE_SIZE / 1024 / 1024}MB limit'
                }
            file_ext = Path(filename).suffix.lower()
            logger.info(f"Detected extension: {file_ext} for file: {filename}")
            file_type = self.SUPPORTED_EXTENSIONS.get(file_ext)
            if not file_type:
                return {
                    'text': None,
                    'images': [],
                    'file_type': file_ext,
                    'error': f'Unsupported file type: {file_ext}. Supported types: PDF, DOCX, TXT'
                }
            if file_type == 'pdf':
                text = self._extract_pdf(file_content)
                images = self._extract_pdf_images(file_content)
            elif file_type == 'docx':
                text = self._extract_docx(file_content)
                images = []
            elif file_type == 'txt':
                text = self._extract_txt(file_content)
                images = []
            else:
                raise ValueError(f"Unhandled file type: {file_type}")
            if not text or not text.strip():
                return {
                    'text': None,
                    'images': images,
                    'file_type': file_type,
                    'error': 'No text content found in file'
                }
            logger.info(f"Successfully extracted {len(text)} characters and {len(images)} images from {filename}")
            return {
                'text': text.strip(),
                'images': images,
                'file_type': file_type,
                'error': None
            }
        except Exception as e:
            logger.error(f"Error extracting text/images from {filename}: {str(e)}", exc_info=True)
            return {
                'text': None,
                'images': [],
                'file_type': None,
                'error': f'Failed to extract text/images: {str(e)}'
            }

    def _extract_pdf_images(self, file_content: bytes):
        """Extract images from PDF file using pypdf's page.images API.
        
        Uses the modern page.images property which automatically recurses
        into Form XObjects — handles design-tool PDFs where images are nested.
        Returns a list of dicts with image bytes and metadata.
        """
        images = []
        try:
            reader = PdfReader(BytesIO(file_content))
            for page_num, page in enumerate(reader.pages):
                try:
                    for img in page.images:
                        try:
                            data = img.data
                            # Determine format from the image name extension
                            name = img.name or ""
                            if name.lower().endswith(".jp2"):
                                fmt = "jp2"
                            elif name.lower().endswith((".jpg", ".jpeg")):
                                fmt = "jpeg"
                            elif name.lower().endswith(".png"):
                                fmt = "png"
                            else:
                                fmt = "jpeg"  # default — most PDF images are JPEG
                            # Get dimensions via PIL (page.images doesn't expose w/h directly)
                            try:
                                pil_img = Image.open(BytesIO(data))
                                width, height = pil_img.size
                            except Exception:
                                width, height = None, None
                            images.append({
                                'data': data,
                                'format': fmt,
                                'width': width,
                                'height': height,
                                'page': page_num + 1
                            })
                        except Exception as e:
                            logger.warning(f"Failed to extract image '{img.name}' from page {page_num + 1}: {e}")
                except Exception as e:
                    logger.warning(f"Failed to iterate images on page {page_num + 1}: {e}")
            logger.info(f"PDF image extraction: found {len(images)} image(s) across {len(reader.pages)} pages")
            return images
        except Exception as e:
            logger.warning(f"PDF image extraction failed: {e}")
            return []
        """
        Extract text from file content.
        
        Args:
            file_content: Raw file bytes
            filename: Original filename
            
        Returns:
            dict with keys:
                - text: Extracted text content
                - file_type: Detected file type
                - error: Error message if extraction failed
        """
        try:
            # Check file size
            if len(file_content) > self.MAX_FILE_SIZE:
                return {
                    'text': None,
                    'file_type': None,
                    'error': f'File size exceeds {self.MAX_FILE_SIZE / 1024 / 1024}MB limit'
                }
            
            # Detect file type from extension
            file_ext = Path(filename).suffix.lower()
            logger.info(f"Detected extension: {file_ext} for file: {filename}")
            
            # Get file type
            file_type = self.SUPPORTED_EXTENSIONS.get(file_ext)
            if not file_type:
                return {
                    'text': None,
                    'file_type': file_ext,
                    'error': f'Unsupported file type: {file_ext}. Supported types: PDF, DOCX, TXT'
                }
            
            # Extract based on type
            if file_type == 'pdf':
                text = self._extract_pdf(file_content)
            elif file_type == 'docx':
                text = self._extract_docx(file_content)
            elif file_type == 'txt':
                text = self._extract_txt(file_content)
            else:
                raise ValueError(f"Unhandled file type: {file_type}")
            
            # Validate extracted text
            if not text or not text.strip():
                return {
                    'text': None,
                    'file_type': file_type,
                    'error': 'No text content found in file'
                }
            
            logger.info(f"Successfully extracted {len(text)} characters from {filename}")
            
            return {
                'text': text.strip(),
                'file_type': file_type,
                'error': None
            }
            
        except Exception as e:
            logger.error(f"Error extracting text from {filename}: {str(e)}", exc_info=True)
            return {
                'text': None,
                'file_type': None,
                'error': f'Failed to extract text: {str(e)}'
            }
    
    def _extract_pdf(self, file_content: bytes) -> str:
        """Extract text from PDF file."""
        try:
            reader = PdfReader(BytesIO(file_content))
            text_parts = []
            
            for page_num, page in enumerate(reader.pages):
                try:
                    text = page.extract_text()
                    if text:
                        text_parts.append(text)
                except Exception as e:
                    logger.warning(f"Failed to extract text from PDF page {page_num}: {e}")
                    continue
            
            return '\n\n'.join(text_parts)
            
        except Exception as e:
            raise ValueError(f"PDF extraction failed: {str(e)}")
    
    def _extract_docx(self, file_content: bytes) -> str:
        """Extract text from DOCX file."""
        try:
            doc = Document(BytesIO(file_content))
            text_parts = []
            
            # Extract from paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            
            # Extract from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = ' | '.join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        text_parts.append(row_text)
            
            return '\n\n'.join(text_parts)
            
        except Exception as e:
            raise ValueError(f"DOCX extraction failed: {str(e)}")
    
    def _extract_txt(self, file_content: bytes) -> str:
        """Extract text from TXT file."""
        try:
            # Try UTF-8 first
            try:
                return file_content.decode('utf-8')
            except UnicodeDecodeError:
                # Fallback to latin-1
                return file_content.decode('latin-1')
                
        except Exception as e:
            raise ValueError(f"TXT extraction failed: {str(e)}")


# Singleton instance
file_extractor = FileExtractor()
